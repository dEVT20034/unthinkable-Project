import os, math, shutil
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import requests

# ========== Config ==========
PORT = int(os.environ.get("PORT", 5501))

# Load .env only in local dev; on Render, env vars are injected
try:
    from dotenv import load_dotenv  # optional in prod
    load_dotenv()
except Exception:
    pass

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY is not set. Add it in your host's Environment.")

MODELS_LIST_ENDPOINT = "https://generativelanguage.googleapis.com/v1/models"
EMB_MODEL = "text-embedding-004"
EMB_ENDPOINT = f"https://generativelanguage.googleapis.com/v1/models/{EMB_MODEL}:embedContent"

ALLOWED_EXT = {".pdf", ".docx", ".txt"}
MAX_FILES = 10
MAX_FILE_MB = 20

app = Flask(__name__, static_folder="public", static_url_path="/")
CORS(app)
UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ========== In-memory KB ==========
KB = []         # {id, source, text, embedding}
NEXT_ID = 1

# ========== Dynamic model selection ==========
GEN_MODEL = None
GEN_ENDPOINT = None


def _list_models():
    r = requests.get(
        MODELS_LIST_ENDPOINT,
        headers={"x-goog-api-key": GEMINI_API_KEY},
        timeout=60
    )
    if not r.ok:
        print("[LIST MODELS ERROR]", r.status_code, r.text[:600])
        r.raise_for_status()
    return r.json().get("models", [])


def _supports_generate(m):
    methods = m.get("supportedGenerationMethods") or m.get("generationMethods") or []
    methods_lower = [s.lower() for s in methods]
    return "generatecontent" in methods_lower


def _choose_generation_model():
    global GEN_MODEL, GEN_ENDPOINT
    if GEN_MODEL and GEN_ENDPOINT:
        return GEN_MODEL, GEN_ENDPOINT

    models = _list_models()
    by_name = {m["name"]: m for m in models if "name" in m}

    preferred = [
        "models/gemini-1.5-flash",
        "models/gemini-1.5-flash-001",
        "models/gemini-1.5-flash-8b",
        "models/gemini-1.5-pro",
        "models/gemini-1.5-pro-001",
        "models/gemini-1.0-pro",
    ]
    for name in preferred:
        m = by_name.get(name)
        if m and _supports_generate(m):
            GEN_MODEL = name.split("/", 1)[1]
            GEN_ENDPOINT = f"https://generativelanguage.googleapis.com/v1/models/{GEN_MODEL}:generateContent"
            print("[GEN MODEL] selected:", GEN_MODEL)
            return GEN_MODEL, GEN_ENDPOINT

    for m in models:
        if _supports_generate(m):
            GEN_MODEL = m["name"].split("/", 1)[1]
            GEN_ENDPOINT = f"https://generativelanguage.googleapis.com/v1/models/{GEN_MODEL}:generateContent"
            print("[GEN MODEL] fallback selected:", GEN_MODEL)
            return GEN_MODEL, GEN_ENDPOINT

    raise RuntimeError("No available models support generateContent for this API key/tenant.")


# ========== Helpers ==========
def ok(**data):
    return jsonify({"ok": True, **data})


def err(msg, code=400):
    return jsonify({"ok": False, "error": msg}), code


def allowed_file(name: str) -> bool:
    return os.path.splitext(name or "")[1].lower() in ALLOWED_EXT


def sliding(text: str, size: int = 1000, overlap: int = 120):
    n = len(text)
    if n == 0:
        return
    step = max(1, size - overlap)
    i = 0
    while i < n:
        end = min(n, i + size)
        chunk = text[i:end]
        last_dot = chunk.rfind(". ")
        if last_dot > size * 0.6 and end < n:
            chunk = chunk[: last_dot + 1]
        t = chunk.strip()
        if t:
            yield t
        i += step


def cosine(a, b):
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a)) or 1e-12
    nb = math.sqrt(sum(y * y for y in b)) or 1e-12
    return dot / (na * nb)


# ----------- Enhanced DOCX extractor (paragraphs, tables, headers, footers) -----------
def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = []

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)

    for section in doc.sections:
        try:
            for p in section.header.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
        except Exception:
            pass
        try:
            for p in section.footer.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
        except Exception:
            pass

    cleaned = [s.strip() for s in parts if s and s.strip()]
    seen, uniq = set(), []
    for s in cleaned:
        if s not in seen:
            seen.add(s)
            uniq.append(s)
    return "\n".join(uniq)


def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".docx":
        return read_docx_text(filepath)
    if ext == ".pdf":
        return pdf_extract_text(filepath) or ""
    raise ValueError(f"Unsupported file type: {ext}")


def gemini_embed(text: str):
    body = {"content": {"parts": [{"text": text}]}}
    r = requests.post(
        EMB_ENDPOINT,
        json=body,
        headers={"x-goog-api-key": GEMINI_API_KEY, "Content-Type": "application/json"},
        timeout=90,
    )
    if not r.ok:
        print("[EMBED ERROR]", r.status_code, r.text[:600])
        r.raise_for_status()
    jd = r.json()
    vec = (
        jd.get("embedding", {}).get("values")
        or jd.get("embedding", {}).get("value")
        or jd.get("embedding", [])
    )
    if not vec:
        print("[EMBED WARN] Empty embedding payload:", jd)
    return vec


def add_doc_to_kb(filepath: str, original_name: str) -> int:
    global NEXT_ID
    text = extract_text(filepath)
    joined = "\n".join(s.strip() for s in text.splitlines() if s.strip())
    if not joined:
        raise ValueError(
            f"No extractable text found in '{original_name}'. "
            f"If this is a scanned document, convert to text or enable OCR."
        )
    added = 0
    for ch in sliding(joined, size=1000, overlap=120):
        vec = gemini_embed(ch)
        KB.append({"id": NEXT_ID, "source": original_name, "text": ch, "embedding": vec})
        NEXT_ID += 1
        added += 1
    return added


# Always return JSON for errors (prevents HTML error pages)
@app.errorhandler(Exception)
def handle_any_exception(e):
    code = getattr(e, "code", 500)
    if isinstance(e, RuntimeError) and isinstance(getattr(e, "args", [None])[0], dict):
        return jsonify(ok=False, **e.args[0]), 500
    return jsonify(ok=False, error=str(e)), code


# ========== Routes ==========
@app.route("/")
def index():
    return send_from_directory("public", "index.html")


@app.route("/health", methods=["GET"])
def health():
    model_name = GEN_MODEL
    try:
        if not model_name:
            model_name, _ = _choose_generation_model()
    except Exception:
        pass
    return ok(chunks=len(KB), model=model_name or "—")


@app.route("/models", methods=["GET"])
def models():
    try:
        ms = _list_models()
        return ok(models=ms[:100])
    except Exception as e:
        return err(str(e), 500)


@app.route("/gemini/ping", methods=["GET"])
def gemini_ping():
    try:
        model, endpoint = _choose_generation_model()
        r = requests.post(
            endpoint,
            json={"contents": [{"role": "user", "parts": [{"text": "Say 'pong'."}]}]},
            headers={"x-goog-api-key": GEMINI_API_KEY, "Content-Type": "application/json"},
            timeout=60,
        )
        return jsonify(ok=r.ok, model=model, sample=r.text[:300])
    except Exception as e:
        return jsonify(ok=False, error=str(e)), 500


@app.route("/reset", methods=["POST"])
def reset():
    global KB, NEXT_ID
    KB, NEXT_ID = [], 1
    removed = 0
    try:
        if os.path.isdir(UPLOAD_DIR):
            for _, _, files in os.walk(UPLOAD_DIR):
                removed += len(files)
            shutil.rmtree(UPLOAD_DIR, ignore_errors=True)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
    except Exception as e:
        return ok(message=f"KB cleared, but cleanup warning: {e}", chunks=0, filesRemoved=removed)
    return ok(message="Knowledge base and uploads cleared.", chunks=0, filesRemoved=removed)


@app.route("/upload", methods=["POST"])
def upload():
    files = request.files.getlist("files")
    if not files:
        return err("No files uploaded")
    if len(files) > MAX_FILES:
        return err(f"Too many files (max {MAX_FILES})")

    results = []
    for f in files:
        name = f.filename or ""
        if not name:
            continue
        if not allowed_file(name):
            return err("Only PDF, DOCX, and TXT are allowed")
        try:
            f.stream.seek(0, os.SEEK_END)
            size = f.stream.tell()
            f.stream.seek(0)
            if size > MAX_FILE_MB * 1024 * 1024:
                return err(f"{name}: exceeds {MAX_FILE_MB} MB")
        except Exception:
            pass

        safe = secure_filename(name)
        path = os.path.join(UPLOAD_DIR, safe)
        f.save(path)

        try:
            added = add_doc_to_kb(path, safe)
            results.append({"file": safe, "chunksAdded": added})
        except Exception as e:
            results.append({"file": safe, "error": str(e)})

    total_chunks = len(KB)
    if not any(r.get("chunksAdded", 0) > 0 for r in results):
        return err({"message": "No text indexed from uploads", "details": results})

    return ok(results=results, chunks=total_chunks)


@app.route("/suggest", methods=["GET"])
def suggest():
    defaults = [
        "Summarize the document briefly.",
        "What is the project name and its owner?",
        "List the key requirements.",
        "What are the deadlines and milestones?",
        "What risks or blockers are mentioned?",
    ]
    if not KB:
        return ok(suggestions=defaults)

    sources = {}
    for ch in KB[:10]:
        sources[ch["source"]] = sources.get(ch["source"], 0) + 1
    top_sources = [s for s, _ in sorted(sources.items(), key=lambda x: -x[1])[:2]]

    text_sample = " ".join(ch["text"] for ch in KB[:5])[:5000]
    words = [w.strip(",.():;") for w in text_sample.split()]
    caps = [w for w in words if len(w) > 2 and w[0].isupper() and w.isalpha()]
    freq = {}
    for w in caps:
        freq[w] = freq.get(w, 0) + 1
    top_terms = [w for w, _ in sorted(freq.items(), key=lambda x: -x[1])[:5]]

    sug = []
    for s in top_sources:
        sug.extend([
            f"Give a concise summary of {s}.",
            f"What are the key requirements in {s}?",
            f"Who are the stakeholders mentioned in {s}?",
        ])
    if top_terms:
        sug.extend([
            f"What does {top_terms[0]} refer to in the documents?",
            f"How is {top_terms[0]} related to {top_terms[1] if len(top_terms)>1 else 'the project'}?",
        ])

    seen, out = set(), []
    for q in sug + defaults:
        if q not in seen:
            seen.add(q)
            out.append(q)
        if len(out) >= 8:
            break

    return ok(suggestions=out)


@app.route("/ask", methods=["POST"])
def ask():
    data = request.get_json(silent=True) or {}
    question = (data.get("question") or "").strip()
    topK = max(1, min(int(data.get("topK", 5)), 20))

    if not question:
        return err("Missing question")
    if not KB:
        return err("No knowledge uploaded yet")

    model, endpoint = _choose_generation_model()

    qv = gemini_embed(question)
    scored = sorted(
        ({"score": cosine(qv, ch["embedding"]), "ch": ch} for ch in KB),
        key=lambda x: x["score"],
        reverse=True,
    )[:topK]

    context = "\n\n".join(f'[{s["ch"]["id"]} | {s["ch"]["source"]}] {s["ch"]["text"]}' for s in scored)

    prompt = (
        "You are a helpful assistant. Answer ONLY using the provided context. "
        "If the answer is not present, say you don't have enough information.\n\n"
        f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
    )

    r = requests.post(
        endpoint,
        json={"contents": [{"role": "user", "parts": [{"text": prompt}]}]},
        headers={"x-goog-api-key": GEMINI_API_KEY, "Content-Type": "application/json"},
        timeout=90,
    )
    if not r.ok:
        print("[GEN ERROR]", r.status_code, r.text[:800])
        r.raise_for_status()

    jd = r.json()
    answer = "\n".join(
        (p.get("text", "") or "")
        for p in jd.get("candidates", [{}])[0].get("content", {}).get("parts", [])
    ).strip()

    return ok(
        answer=answer,
        citations=[{"id": s["ch"]["id"], "source": s["ch"]["source"], "score": round(s["score"], 3)} for s in scored],
        model=model,
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT, debug=bool(int(os.environ.get("FLASK_DEBUG", "0"))))